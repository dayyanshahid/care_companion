"""Knowledge base for the Care Companion FAQ (the RAG layer).

Ingestion parses the FAQ .docx into Q&A chunks, embeds each with OpenAI
(text-embedding-3-small by default), and stores them in MongoDB - the text
and its vector on the same document, in `faq_chunks`.

Retrieval embeds the query and scores every chunk by cosine similarity here
in Python. This MongoDB has no vector index, and at a few dozen chunks a
full scan is faster than reaching for one would be.
"""
import re
from math import sqrt

import openai
from django.conf import settings

from database.models import FaqChunk
from database.serializers import FaqChunkSerializer
from utils.common import build_error
from utils.enums import HttpStatus
from utils.messages import messages, KnowledgeError

_QA_RE = re.compile(r"^Q\d+\.\s*", re.IGNORECASE)
_SECTION_RE = re.compile(r"^(\d+)\.\s+(.*)")

_openai_client = None


# --- Ingestion --------------------------------------------------------------

def ingest_faq(path):
    """Parse the FAQ .docx, embed each Q&A, and replace the stored chunks."""
    entries = parse_faq(path)

    if not entries:
        raise KnowledgeError(messages["noFaqEntries"].format(path=path))

    vectors = _embed(
        [f"{e['category']} — {e['question']}\n{e['answer']}" for e in entries]
    )

    # Everything is replaced in one go, so a re-ingest cannot leave a stale
    # answer behind next to the new ones.
    FaqChunk.objects.all().delete()
    FaqChunk.objects.bulk_create(
        [
            FaqChunk(embedding=vector, **entry)
            for entry, vector in zip(entries, vectors)
        ]
    )

    return len(entries)


def parse_faq(path):
    """Extract Q&A entries from the FAQ .docx.

    The FAQ stores each section header and each question in its own table.
    A section cell looks like "3. Financial Coverage..." followed by
    "5 questions"; a question cell starts with "Qn.".
    """
    from docx import Document  # imported here so parsing is optional at runtime

    document = Document(path)
    entries = []
    category = "General"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                if not paragraphs:
                    continue

                head = paragraphs[0]
                section = _SECTION_RE.match(head)
                if section and "question" in paragraphs[-1].lower():
                    category = section.group(2).strip()
                    continue

                if _QA_RE.match(head):
                    question = _QA_RE.sub("", head).strip()
                    answer = " ".join(paragraphs[1:]).strip()
                    if question and answer:
                        entries.append(
                            {"category": category, "question": question, "answer": answer}
                        )
    return entries


# --- Retrieval --------------------------------------------------------------

def search(query, top_k=None):
    """The search endpoint's whole job: the ranked chunks, ready to return."""
    try:
        results = retrieve(query, top_k=top_k or settings.RAG_TOP_K)
    except KnowledgeError as exc:
        raise build_error(
            messages["searchUnavailable"], HttpStatus.badGateway, exc
        ) from exc

    data = []

    for chunk, score in results:
        item = FaqChunkSerializer(chunk).data
        item["score"] = round(score, 4)
        data.append(item)

    return data


def retrieve(query, top_k):
    """Return the top_k (chunk, score) pairs most relevant to the query."""
    chunks = all_chunks()

    if not chunks:
        raise KnowledgeError(messages["noFaqStored"])

    vector = _embed([query])[0]

    ranked = sorted(
        ((chunk, _cosine(vector, chunk.embedding)) for chunk in chunks),
        key=lambda pair: pair[1],
        reverse=True,
    )

    return ranked[:top_k]


def all_chunks(limit=1000):
    """Every stored chunk. For retrieval, and for checks and inspection."""
    return list(FaqChunk.objects.all()[:limit])


def count_chunks():
    """How many chunks are stored. 0 also means nothing has been ingested."""
    return FaqChunk.objects.count()


def _cosine(left, right):
    """Similarity of two vectors, 0 when either has no length."""
    if not right:
        return 0.0

    dot = sum(a * b for a, b in zip(left, right))
    size = sqrt(sum(a * a for a in left)) * sqrt(sum(b * b for b in right))

    return dot / size if size else 0.0


# --- Clients ----------------------------------------------------------------

def _embed(texts):
    try:
        response = _client().embeddings.create(input=texts, model=settings.EMBEDDING_MODEL)
        return [item.embedding for item in response.data]
    except Exception as exc:  # network, auth, rate limit, etc.
        raise KnowledgeError(str(exc)) from exc


def _client():
    global _openai_client
    if _openai_client is None:
        if not settings.OPENAI_API_KEY:
            raise KnowledgeError(messages["openaiKeyMissing"])
        _openai_client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
    return _openai_client
